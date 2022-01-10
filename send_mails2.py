import xlrd
import os
import json
import smtplib
# from email.message import EmailMessage
from email.mime.text import MIMEText

from email.header import Header
from email.mime.base import MIMEBase
from email import encoders
from email.mime.multipart import MIMEMultipart

from docx import Document

from pathlib import Path

import re
from random import randint

# prerequisites:
DEBUG = True
DEBUG_TO = ["lesoleil@ukr.net", "viktorborodinmm@gmail.com"]  # To whomever you want to send the mail

DIRSTUD = "~/Лекции/Students/"
FILE_JSON = "2021_1.json"

KURS = "Java"  # "C++" # "Python"
NUM_STUD = 35
ADD_TASKS = False
GROUP_NAME = "All"  # "ИВт" #
RND_QUESTIONS = True  # False #

LECT = True  # False
MODE = 2  # start  # 1 # Homework # 2 # start exam 3 # end exam

EXAM_PAPERS = "2021_java_papers.docx"  # "2021_cpp_papers.docx"  # "2021_cpp_papers.docx" #
# column for papers numbers, if None then random
Nep = None
PAPER_NAME = "2021_java_ep_{}.docx"  # "2021_cpp_ep_{}.docx"

# TODO: XML
# data JSON

students_data = {

    "Java": {
        "title": "Мова програмування Java",
        "man_books": ['LabsJava1_2.pdf'],
        "table": "2021_Java_3kurs.xlsx",
        "repa": "https://github.com/mecchmatProjects/JavaCourse",
        "telega": "",
        "doc": "https://docs.google.com/spreadsheets/d/"
               "1mydlx2eVudSBAacIzTn8v7A2Z8Rdqc2Qf-Y15pWLzEs/edit?usp=sharing",
        "groups": [
            {
                "page": "KM1",
                "repa": "https://github.com/mecchmatProjects/JavaPracticum2",
                "doc": "https://docs.google.com/spreadsheets/d/"
                       "1rUHfT795-GxrSvZFziE404El7dvR-by9c1pSQTfWF4I/edit?usp=sharing"

            },
            {
                "page": "KM2",
                "repa": "https://github.com/mecchmatProjects/JavaPracticum2",
                "doc": "https://docs.google.com/spreadsheets/d/"
                       "1rUHfT795-GxrSvZFziE404El7dvR-by9c1pSQTfWF4I/edit?usp=sharing"

            },
            {

            }

        ]
    },

    "C++": {
        "title": "Мова програмування C++",
        "man_books": ["methodicalCpp.pdf"],  # "lectC.pdf"
        "table": "2021_Cpp_2_1kurs.xlsx",
        "repa": "https://github.com/mecchmatProjects/CppCourse",
        "telega": "https://t.me/joinchat/Hz8gY26Sa9U3YzE6",
        "doc": "https://docs.google.com/spreadsheets/d/"
               "1t7UudixlMkDeknPluA2rRUFMUMjCQoVILtYvMgwDChg/edit?usp=sharing",
        "groups": [
            {
                "page": "KM1",
                "repa": "https://github.com/mecchmatProjects/",
                "doc": "https://docs.google.com/spreadsheets/d/"
                       "1t7UudixlMkDeknPluA2rRUFMUMjCQoVILtYvMgwDChg/edit?usp=sharing"

            },
            {
                "page": "KM2",
            },
            {

            }

        ]
    },

}

# students_data = json.load(FILE_JSON)
STUD_TABLE = os.path.join(DIRSTUD, students_data[KURS]["table"])  # "2021_Cpp_2_1kurs.xlsx" #"2021_Java_KM1.xlsx"

print(STUD_TABLE)


wb = xlrd.open_workbook(STUD_TABLE)
sheet_names = wb.sheet_names()
print(sheet_names)

# base 0,1
lst_to = []
lst_names = []

lst_hw = []
Hw_Columns = []  # [i for i in range(19,30)]

# ex_papers

lst_exam_papers = []
# marks
lst_marks = []

if MODE == 0:

    lst_sheet_name = [sheet_names[0]]

else:
    # homeworks
    lst_sheet_name = []  # [item for item in students_data[KURS]["groups"]]

    if GROUP_NAME == "All":
        for x in students_data[KURS]["groups"]:
            if x:
                print("Page name", x['page'])
                lst_sheet_name.append(x['page'])
    else:
        lst_sheet_name = [GROUP_NAME]

if MODE > 0:
    for sheet_name in lst_sheet_name:

        sh = wb.sheet_by_name(sheet_name)
        print(sh.nrows, sh.ncols)
        Nmark = sh.ncols - 1
        for i in range(2, sh.nrows):
            fio = sh.col_values(1)[i]
            if not fio:
                continue
            address = sh.col_values(2)[i]

            hws = ""
            # print(Hw_Columns)
            for col in Hw_Columns:
                hw = sh.col_values(col)[i]
                hws += str(hw) + ";\t"
                # print(hw)

            paper = int(sh.col_values(Nep)[i]) if Nep else 0
            print("p", paper)

            mark = sh.col_values(Nmark)[i]

            lst_to.append(address)
            lst_names.append(fio)
            lst_exam_papers.append(paper)
            lst_hw.append(hws)
            lst_marks.append(mark)
# just debug:
k = 5
for a, b in zip(lst_to[-k:], lst_names[-k:]):
    print(a, b)

for a, b in zip(lst_exam_papers[-k:], lst_hw[-k:]):
    print(a, b)
input()

# Exam Preparation
PATTERN_EX_BETW = r"Екзамінаційний білет № \d+(.*?)\n"
PATTERN_EX = r"Екзамінаційний білет № (\d+)"
PATTERN_DIG = r"(\d+)."
PATTERN_PAR = r"^(\d+).\t"

pat1 = re.compile(re.escape(PATTERN_EX), re.IGNORECASE)

doc_papers = Document(EXAM_PAPERS)

dict_papers = {}
k = 0
question_indx = 0
tmp = " "
found_header = False

total_text = []
question = ""
for p in doc_papers.paragraphs:

    for r in p.runs:
        if not r.text:
            if question:
                total_text.append(question)
                question = ""
            continue

        # text_search = re.search(PATTERN_EX, r.text)
        # print("paragprahp",text_search)

        question += r.text

        # input()

# print(total_text)
# input("total")

tr = ""
nq = 0
for ticket in total_text:
    num = re.search(PATTERN_EX, ticket)
    if num:
        k = int(num.group(1))
        dict_papers[k - 1] = tr
        nq = 0
        tr = ""
    if nq != 0:
        tr += "\n\n" + str(nq) + "\t" + ticket
    else:
        tr += ticket
    nq += 1
    # print(tr)

if ADD_TASKS:

    ZADACH = "zada4i.docx"
    doc_tasks = Document(ZADACH)
    task = ""
    all_tasks = []
    indx = 0
    for p in doc_tasks.paragraphs:

        for r in p.runs:
            if not r.text:
                if task:
                    # all_tasks.append(task)
                    print(task)
                    # input()
                    dict_papers[indx] += "\n\n3.\t" + task
                    indx += 1
                    task = ""
                continue

            # text_search = re.search(PATTERN_EX, r.text)
            # print("paragprahp",text_search)
            task += r.text

            # input()

for x, y in dict_papers.items():
    print("numer ", x)
    print("text: ", y)
    # write to file
    print(PAPER_NAME.format(x))

    dname = PAPER_NAME.format(x)
    document = Document()
    document.add_paragraph(y)
    document.save(dname)

# email body
# Name,course, repa, repa addres
body0 = """
Відомості по курсу {}

Шановний(шановна) {}!
Висилаю координати курсу "{}":
Репозиторій з матеріалами лекцій:
{}

У вкладенні пересилаю задачник з 4 теми.

Поточні оцінки та домашні завдання Гугл-документі:
{}

З повагою, 
Ваш викладач
Віктор Бородін. 

Не відповідайте на цей лист - тут працює робот.
У випадку питань пишить на viktorborodinmm@gmail.com

Please don't reply: this message is sent by robot.
In case of any issue reply to viktorborodinmm@gmail.com
"""
# name, repa_lect, title, HW, mark, doc
body1 = """
{} Homework 
Шановний(шановна) {}!
Репозиторій з матеріалами лекцій:
{}
У вкладенні пересилаю задачник з 5 теми.

Ваші домашні завдання з курсу "{}" {}

Ваші поточні бали {}.

Докладніше на Google-docs {}.

Віктор Бородін

Не відповідайте на цей лист - тут працює робот.
У випадку питань пишить на viktorborodinmm@gmail.com

Please don't reply: this message is sent by robot.
In case of any issue reply to viktorborodinmm@gmail.com
"""
# name, title, num, text
body_exst = """
Шановний(шановна){}!
Ваш білет з курсу "{}" номер {} у вкладенні.
Текст білету продубльований сюди:
{} 


Якщо погано відображається - то дивиться у вкладення.

У випадку запитань або проблем з відкриттям 
звертайтеся до викладача на конференцію або телеграм. 

Віктор Бородін

Не відповідайте на цей лист - тут працює робот.
У випадку питань пишить на viktorborodinmm@gmail.com

Please don't reply: this message is sent by robot.
In case of any issue reply to viktorborodinmm@gmail.com

"""

body_ex_end = """
Шановний(шановна){}!
Ваша оцінка за іспит з курсу "{}" {} та ітогова оцінка {}.
Якщо ви згодні з цією оцінкою пришлить відповідь на viktorborodinmm@gmail.com
або телеграм.
 
Інакше, чекаю вас на продовженні іспиту в домовлений час.

Віктор Бородін 

Довідкова інформація на Google-docs {}

Не відповідайте на цей лист - тут працює робот.
У випадку питань пишить на viktorborodinmm@gmail.com

Please don't reply: this message is sent by robot.
In case of any issue reply to viktorborodinmm@gmail.com
"""

EMAIL = 'vborodin001@gmail.com'
PASSWORD = 'iaGM7ahwae'
SERVER = 'imap.gmail.com'

email_id = 'vborodin001@gmail.com'
email_pass = 'iaGM7ahwae'

SMTP_HOST = 'smtp.gmail.com'
SMTP_PORT = 465


REPA_LECT = students_data[KURS]["repa"]

COURSE = students_data[KURS]["title"]
REPA = students_data[KURS]["groups"][0]["repa"]
ADDRESS_GD = students_data[KURS]["groups"][0]["doc"]

print(lst_names)

where = DEBUG_TO if DEBUG else lst_to

for k, to in enumerate(where):  # enumerate(DEBUG_TO):  #

    name = lst_names[k]

    # exam start
    # random  or chosen
    num_pap = randint(1, NUM_STUD) if RND_QUESTIONS else lst_exam_papers[k]

    paper_item = dict_papers[num_pap]
    # print(body) # body.format(r1,r2,)
    print("{}: to {} for {} ".format(k, to, name))

    # prapere email
    msg = MIMEMultipart()
    if MODE == 0:
        msg.attach(MIMEText(body0.format(KURS, name, COURSE, REPA, ADDRESS_GD), 'plain', 'utf-8'))
    elif MODE == 1:
        # name, repa_lect, title, HW, mark, doc
        Text = body1.format(KURS, name, students_data[KURS]["repa"], COURSE, lst_hw[k], lst_marks[k], ADDRESS_GD
                            )
        print(Text)

        msg.attach(MIMEText(Text, 'plain', 'utf-8'))
    elif MODE == 2:
        Text = body_exst.format(name, COURSE, num_pap, paper_item)
        print(Text)

        msg.attach(MIMEText(Text, 'plain', 'utf-8'))
    elif MODE == 3:
        Text = body_ex_end.format(name, COURSE, lst_marks, lst_marks, ADDRESS_GD)
        print(Text)

        msg.attach(MIMEText(Text, 'plain', 'utf-8'))

    msg['Subject'] = Header("C++ Course" if KURS == "C++" else 'Java Course', 'utf-8')
    msg['From'] = email_id
    msg['To'] = to  # ", ".join(recipients_emails)

    if MODE == 1:
        files = ['LabsJava5.pdf']  # students_data[KURS]["man_books"] # ,  ['lectС.pdf', 'methodicalCPP.pdf']  #

        for path in files:
            # path = os.path.join("~/Лекции/Cpp/CppCourse","methodicalCPP.pdf")
            part = MIMEBase('application', "octet-stream")
            with open(path, 'rb') as file:
                part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition',
                            'attachment; filename={}'.format(Path(path).name), filename=Path(path).name)
            print("attached: ", Path(path).name)
            msg.attach(part)
    elif MODE == 2:
        files = [PAPER_NAME.format(num_pap)]

        for path in files:
            # path = os.path.join("~/Лекции/Cpp/CppCourse","methodicalCPP.pdf")
            part = MIMEBase('application', "octet-stream")
            with open(path, 'rb') as file:
                part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition',
                            'attachment; filename={}'.format(Path(path).name), filename=Path(path).name)
            print("attached: ", Path(path).name)
            msg.attach(part)

    s = smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT)
    # s =smtplib.SMTP('smtp.gmail.com', 465, timeout=10) # 587
    # s.set_debuglevel(1)
    try:
        # s.starttls()
        s.login(EMAIL, PASSWORD)
        s.sendmail(msg['From'], to, msg.as_string())
    finally:
        s.quit()
        print("All sent")
