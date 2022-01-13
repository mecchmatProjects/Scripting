import docx
import re

doc = docx.Document('T23.3.docx')
for p in range(len(doc.paragraphs)):
    for r in range(len(doc.paragraphs[p].runs)):
        if re.match('^M{0,3}(CM|CD|D?C{0,3})?(XC|XL|L?X{0,3})?(IX|IV|V?I{0,3})?$', doc.paragraphs[p].runs[r].text):
            doc.paragraphs[p].runs[r].font.color.rgb = docx.shared.RGBColor(0, 255, 0)
            doc.paragraphs[p].runs[r].bold = True
            print('green')
        elif re.match('[MCDXLIV]+', doc.paragraphs[p].runs[r].text):
            doc.paragraphs[p].runs[r].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
            doc.paragraphs[p].runs[r].bold = True
            print('red')

doc.save('T23.3_rez.docx')