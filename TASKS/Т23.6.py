import docx
import re

doc = docx.Document('T23.6.docx')
for p in range(len(doc.paragraphs)):
    for r in range(len(doc.paragraphs[p].runs)):
        if re.match('^\.[0-9]+$', doc.paragraphs[p].runs[r].text.strip()):
            doc.paragraphs[p].runs[r].text = '0'+doc.paragraphs[p].runs[r].text
            doc.paragraphs[p].runs[r].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
            doc.paragraphs[p].runs[r].bold = True
            print('red')
        elif re.match('^[0-9]+\.$', doc.paragraphs[p].runs[r].text.strip()):
            doc.paragraphs[p].runs[r].text = doc.paragraphs[p].runs[r].text+'0'
            doc.paragraphs[p].runs[r].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
            doc.paragraphs[p].runs[r].bold = True
            print('red')
        elif re.match('[0-9]+\.[0-9]+', doc.paragraphs[p].runs[r].text.strip()):
            doc.paragraphs[p].runs[r].font.color.rgb = docx.shared.RGBColor(0, 255, 0)
            doc.paragraphs[p].runs[r].bold = True
            print('green')

doc.save('T23.6_rez.docx')