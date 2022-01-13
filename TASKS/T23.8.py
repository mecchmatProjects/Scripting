import re
from docx import Document
import os

def docx_replace_regex(doc, regex, replace):
    regex = re.compile(regex)

    for p in doc.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def get_list(to_check):
    rez = {}
    rez_list = {}
    if to_check not in rez:
        rez[to_check] = {}
    files = os.listdir(to_check)
    for f in files:
        filename = to_check + '/' + f
        if os.path.isdir(filename):
            rez[to_check][f], rez_list_temp = get_list(filename)
            rez_list.update(rez_list_temp)
        else:
            rez[to_check][f] = {'size': os.path.getsize(filename), 'Last modification': os.path.getmtime(filename)}
            rez_list.update({filename: {'size': os.path.getsize(filename), 'Last modification': os.path.getmtime(filename)}})
    return rez, rez_list

_, listdir = get_list('T23.8')
listdir = list(listdir)

for i in range(len(listdir)-1, -1, -1):
    if listdir[i][-5:] != '.docx':
        del listdir[i]


regex = 'me'
replace = 'you'

for filename in listdir:
    doc = Document(filename)
    docx_replace_regex(doc, regex, replace)
    doc.save(filename)