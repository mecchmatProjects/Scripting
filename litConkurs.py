import email
import imaplib
from datetime import datetime,date
from email.header import decode_header
import email.message
import base64
import quopri
from email import policy
import os

from docx import Document
#from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import Pt

import textract
from bs4 import BeautifulSoup as bs

DATE = datetime(2021, 1, 3) #.strftime("%s")
date_format = "{}-{}-{}".format(DATE.day,DATE.strftime("%b"),DATE.year)
print(DATE, date_format, DATE.date().strftime("%d-%b-%y"))

EMAIL = 'vborodin001@gmail.com'
PASSWORD = 'iaGM7ahwae'
SERVER = 'imap.gmail.com'

THEMES = ("Коллекция фантазий", "Колфан",)

DIR_WORD = "WORDS"
N1 = 5000
N2 = 40000

# connect to the server and go to its inbox
mail = imaplib.IMAP4_SSL(SERVER)
mail.login(EMAIL, PASSWORD)
# we choose the inbox but you can select others
#mail.select('inbox')

print("Logged in! Listing messages...");
status, select_data = mail.select('INBOX')
nmessages = select_data[0].decode('utf-8')
print(nmessages)
status, data = mail.search(None, '(SINCE "{}")'.format(date_format))

# for msg_id in data[0].split():
#     msg_id_str = msg_id.decode('utf-8')
#     print("Fetching message {} of {}".format(msg_id_str,
#                                              nmessages))
#     status, msg_data = mail.fetch(msg_id, '(RFC822)')
#     msg_raw = msg_data[0][1]
#     raw_email_string = msg_raw #.decode('utf-8')
#     #print(raw_email_string)
#
#     msg = email.message_from_bytes(msg_raw,
#                                    _class=email.message.EmailMessage,policy=policy.default)
#     print("subj", msg['Subject'])
#
#
#     #z = msg['Subject'].decode("utf-8", "ignore")
#     subject, encoding = decode_header(msg["Subject"])[0]
#     print("enc",subject, encoding)


# we'll search using the ALL criteria to retrieve
# every message inside the inbox
# it will return with its status and a list of ids
#status, data = mail.search(None, 'ALL')

#f'after:{datetime(2020, 1, 3, 1).strftime("%s")} before:{datetime(2020, 1, 3, 10).strftime("%s")}'
#status, data = mail.search(None, '(SINCE "{}")'.format(date_format)) #.format("2020-01-03"))


print("New emails are %d" % len(data))


# the list returned is a list of bytes separated
# by white spaces on this format: [b'1 2 3', b'4 5 6']
# so, to separate it first we create an empty list
mail_ids = []
# then we go through the list splitting its blocks
# of bytes and appending to the mail_ids list
for block in data:
    # the split function called without parameter
    # transforms the text or bytes into a list using
    # as separator the white spaces:
    # b'1 2 3'.split() => [b'1', b'2', b'3']
    mail_ids += block.split()

# now for every id we'll fetch the email
# to extract its content
for i in mail_ids:
    # the fetch function fetch the email given its id
    # and format that you want the message to be
    status, data = mail.fetch(i, '(RFC822)')

    #print(status, data)

    # the content data at the '(RFC822)' format comes on
    # a list with a tuple with header, content, and the closing
    # byte b')'
    for response_part in data:
        # so if its a tuple...
        if isinstance(response_part, tuple):
            # we go for the content at its second element
            # skipping the header at the first and the closing
            # at the third
            message = email.message_from_bytes(response_part[1],policy=policy.default)

            # with the content we can extract the info about
            # who sent the message and its subject
            mail_from = message['from']
            mail_subject = message['subject']
            print(mail_from, mail_subject)
            #print(quopri.decodestring(mail_subject.encode('utf-8')).decode('utf-8'))
            #mail_subject = base64.b64decode(mail_subject[4:]).decode("UTF-8")

            # if mail_subject not in THEMES:
            #     print("not our theme")
            #     continue

            # downloading attachments
            for part in message.walk():
                # this part comes from the snipped I don't understand yet...
                if part.get_content_maintype() == 'multipart':
                    continue
                if part.get('Content-Disposition') is None:
                    continue
                fileName = part.get_filename()
                print("file "+fileName)
                if not os.path.isdir(DIR_WORD):
                    os.makedirs(DIR_WORD)

                ALLOWED = ['docx','doc','txt']

                if bool(fileName):
                    filePath = os.path.join(DIR_WORD, fileName)
                    print(filePath)
                    with open(filePath, 'wb') as fp:

                        fp.write(part.get_payload(decode=True))
                        fp.close()
                        subject = str(message).split("Subject: ", 1)[1].split("\nTo:", 1)[0]

                    print('Downloaded "{file}" from email titled "{subject}".'.format(file=fileName,
                                                                                subject=mail_subject))

                    form_name = os.path.split(filePath)[-1].split(".")[-1]
                    file_pat = os.path.split(filePath)[-1].split(".")[:-1]
                    print(form_name)
                    if form_name not in ALLOWED:
                        print("Not allowed format")
                    elif form_name=="txt":
                        data = fp.read().replace(" ", "")

                        # get the length of the data
                        number_of_characters = len(data)

                        print('Number of characters in text file :', number_of_characters)
                        if number_of_characters <N1:
                            print("too small")
                        elif number_of_characters >N2:
                            print("too big")

                            lines = fp.readlines()
                            doc = Document(file_pat + ".docx")
                            for line in lines:
                                paragraph = doc.add_paragraph(line)
                                paragraph_format = paragraph.paragraph_format
                                run = paragraph.add_run()
                                font = run.font
                                font.name = 'Calibri'
                                font.size = Pt(12)

                        doc.save(file_pat + ".docx")

                    elif form_name== "doc":

                        soup = bs(
                            open(filePath, encoding="ISO-8859-1").read())
                        [s.extract() for s in soup(['style', 'script'])]
                        tmpText = soup.get_text()
                        text = "".join("".join(tmpText.split('\t')).split('\n')).strip()
                        print(text)

                        document = Document(filePath)
                        s=0
                        for p in document.paragraphs:
                            s += len(p.text)
                        print(s)
                        print('Number of characters in word file :', number_of_characters)
                        if s < N1:
                            print("too small")
                        elif s > N2:
                            print("too big")
                        else:
                            print("ok")



            # # then for the text we have a little more work to do
            # # because it can be in plain text or multipart
            # # if its not plain text we need to separate the message
            # # from its annexes to get the text
            # if message.is_multipart():
            #     mail_content = ''
            #
            #     # on multipart we have the text message and
            #     # another things like annex, and html version
            #     # of the message, in that case we loop through
            #     # the email payload
            #     for part in message.get_payload():
            #         # if the content type is text/plain
            #         # we extract it
            #         if part.get_content_type() == 'text/plain':
            #             mail_content += part.get_payload()
            # else:
            #     # if the message isn't multipart, just extract it
            #     mail_content = message.get_payload()

            # and then let's show its result
            #print('From: {}'.format(mail_from))
            #print('Subject: {}'.format(mail_subject))
            #print('Content: {}'.format(mail_content))