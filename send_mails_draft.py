import xlrd
import os
import json
import smtplib
from email.message import EmailMessage
from email.mime.text import MIMEText

import re
from random import randint


DEBUG_TO=["lesoleil@ukr.net", "viktorborodinmm@gmail.com"]  # To whomever you want to send the mail

DIRSTUD = "~/Лекции/Students/"

FILE_JSON = "2021_1.json"
NUM_STUD = 38
# TODO: XML or JSON

#data JSON

students_data ={

    "Java" : {
          "title": "Мова програмування Java",
          "man_books" : ['LabsJava1_2.pdf'],
          "table" : "2021_Java_3kurs.xlsx",
          "repa" : "https://github.com/mecchmatProjects/JavaCourse",
          "telega": "",
          "doc": "https://docs.google.com/spreadsheets/d/1mydlx2eVudSBAacIzTn8v7A2Z8Rdqc2Qf-Y15pWLzEs/edit?usp=sharing",
          "groups": [
              {
                          "page":"KM1",
                          "repa":"https://github.com/mecchmatProjects/JavaPracticum2",
                          "doc": "https://docs.google.com/spreadsheets/d/1rUHfT795-GxrSvZFziE404El7dvR-by9c1pSQTfWF4I/edit?usp=sharing"

              },
              {
                          "page": "KM2",
                          "repa": "https://github.com/mecchmatProjects/JavaPracticum2",
                          "doc": "https://docs.google.com/spreadsheets/d/1rUHfT795-GxrSvZFziE404El7dvR-by9c1pSQTfWF4I/edit?usp=sharing"

              },
              {

              }

                     ]
    },

   "C++":{
          "title": "Мова програмування C++",
          "man_books" : ["methodicalCpp.pdf"], # "lectC.pdf"
          "table" : "2021_Cpp_2_1kurs.xlsx",
          "repa" : "https://github.com/mecchmatProjects/CppCourse",
          "telega": "https://t.me/joinchat/Hz8gY26Sa9U3YzE6",
          "doc": "https://docs.google.com/spreadsheets/d/1t7UudixlMkDeknPluA2rRUFMUMjCQoVILtYvMgwDChg/edit?usp=sharing",
          "groups": [
              {
                          "page":"KM1",
                          "repa":"https://github.com/mecchmatProjects/",
                          "doc": "https://docs.google.com/spreadsheets/d/1t7UudixlMkDeknPluA2rRUFMUMjCQoVILtYvMgwDChg/edit?usp=sharing"

              },
              {
              },
              {

              }


                     ]
    },


}

# students_data = json.load(FILE_JSON)


KURS = "C++" # "Java" #   "Python"
STUD_TABLE = os.path.join(DIRSTUD,students_data[KURS]["table"])  # "2021_Cpp_2_1kurs.xlsx" #"2021_Java_KM1.xlsx"

print(STUD_TABLE)

LECT = True # False

MODE = 2 # start  # 1 # Homework # 2 # start exam 3 # end exam

# java
ADDRESS_GD_lect = "https://docs.google.com/spreadsheets/d/1mydlx2eVudSBAacIzTn8v7A2Z8Rdqc2Qf-Y15pWLzEs/edit?usp=sharing"
ADDRESS_GD_lab = "https://docs.google.com/spreadsheets/d/1rUHfT795-GxrSvZFziE404El7dvR-by9c1pSQTfWF4I/edit?usp=sharing"

REPA_JAVA = "https://github.com/mecchmatProjects/JavaCourse"

# c++
ADDRESS_CPP_LABS = "https://docs.google.com/spreadsheets/d/1t7UudixlMkDeknPluA2rRUFMUMjCQoVILtYvMgwDChg/edit?usp=sharing"

REPA_CPP = "https://github.com/mecchmatProjects/CppCourse"
TELEGA_CPP = "https://t.me/joinchat/Hz8gY26Sa9U3YzE6"

ADRESS_PLUS = """{}.
               Телеграм канал: {}
              """.format(ADDRESS_CPP_LABS, TELEGA_CPP)

"""
if KURS == "Java":
    STUD_TABLE = "2021_Java_3kurs.xlsx"  # "2021_Java_KM1.xlsx"
    COURSE = "Мова програмування Java" #
    ADDRESS_GD = ADDRESS_GD_lect if LECT else ADDRESS_GD_lab
    REPA = REPA_JAVA

elif KURS == "C++":
    STUD_TABLE = "2021_Cpp_2_1kurs.xlsx"
    COURSE = "Мова програмування C++"  # "Мова програмування Java" #

    # choose data
    ADDRESS_GD = ADRESS_PLUS  # ADDRESS_CPP_LABS  # ADDRESS_GD_lect
    REPA = REPA_CPP  # REPA_JAVA
"""

wb = xlrd.open_workbook(STUD_TABLE)
sheet_names = wb.sheet_names()
print(sheet_names)

# base 0,1
lst_to = []
lst_names = []

lst_hw = []
Hw_Columns = [i for i in range(19,30)]

# ex_papers
Nep = 5
lst_exam_papers = []
# marks
lst_marks = []

if MODE==0:

    lst_sheet_name = [sheet_names[0]]

else:
    # homeworks
    lst_sheet_name = [] #[item for item in students_data[KURS]["groups"]]
    for x in students_data[KURS]["groups"]:
        if x:
            print("Page name",x['page'])
            lst_sheet_name.append(x['page'])

for sheet_name in lst_sheet_name:

    sh = wb.sheet_by_name(sheet_name)
    print(sh.nrows, sh.ncols)
    Nmark = sh.ncols-1
    for i in range(2,sh.nrows):
        fio = sh.col_values(1)[i]
        if not fio:
            continue
        address = sh.col_values(2)[i]

        hws = ""
        print(Hw_Columns)
        for col in Hw_Columns:
            hw = sh.col_values(col)[i]
            hws += str(hw) + ";\t"
            print(hw)

        paper = sh.col_values(Nep)[i]

        mark = sh.col_values(Nmark)[i]

        lst_to.append(address)
        lst_names.append(fio)
        lst_exam_papers.append(paper)
        lst_hw.append(hws)
        lst_marks.append(mark)

k=5
for a,b in zip(lst_to[-k:], lst_names[-k:]):
    print(a,b)

for a,b in zip(lst_exam_papers[-k:], lst_hw[-k:]):
    print(a,b)


#Exam Preparation


from docx import Document
import copy
import itertools

def isolate_run(paragraph, start, end):
    """Return docx.text.Run object containing only `paragraph.text[start:end]`.

    Runs are split as required to produce a new run at the `start` that ends at `end`.
    Runs are unchanged if the indicated range of text already occupies its own run. The
    resulting run object is returned.

    `start` and `end` are as in Python slice notation. For example, the first three
    characters of the paragraph have (start, end) of (0, 3). `end` is not the index of
    the last character. These correspond to `match.start()` and `match.end()` of a regex
    match object and `s[start:end]` of Python slice notation.
    """
    rs = tuple(paragraph._p.r_lst)

    def advance_to_run_containing_start(start, end):
        """Return (r_idx, start, end) triple indicating start run and adjusted offsets.

        The start run is the run the `start` offset occurs in. The returned `start` and
        `end` values are adjusted to be relative to the start of `r_idx`.
        """
        # --- add 0 at end so `r_ends[-1] == 0` ---
        r_ends = tuple(itertools.accumulate(len(r.text) for r in rs)) + (0,)
        r_idx = 0
        while start >= r_ends[r_idx]:
            r_idx += 1
        skipped_rs_offset = r_ends[r_idx - 1]
        return rs[r_idx], r_idx, start - skipped_rs_offset, end - skipped_rs_offset

    def split_off_prefix(r, start, end):
        """Return adjusted `end` after splitting prefix off into separate run.

        Does nothing if `r` is already the start of the isolated run.
        """
        if start > 0:
            prefix_r = copy.deepcopy(r)
            r.addprevious(prefix_r)
            r.text = r.text[start:]
            prefix_r.text = prefix_r.text[:start]
        return end - start

    def split_off_suffix(r, end):
        """Split `r` at `end` such that suffix is in separate following run."""
        suffix_r = copy.deepcopy(r)
        r.addnext(suffix_r)
        r.text = r.text[:end]
        suffix_r.text = suffix_r.text[end:]

    def lengthen_run(r, r_idx, end):
        """Add prefixes of following runs to `r` until `end` is reached."""
        while len(r.text) < end:
            suffix_len_reqd = end - len(r.text)
            r_idx += 1
            next_r = rs[r_idx]
            if len(next_r.text) <= suffix_len_reqd:
                # --- subsume next run ---
                r.text = r.text + next_r.text
                next_r.getparent().remove(next_r)
                continue
            if len(next_r.text) > suffix_len_reqd:
                # --- take prefix from next run ---
                r.text = r.text + next_r.text[:suffix_len_reqd]
                next_r.text = next_r.text[suffix_len_reqd:]

    r, r_idx, start, end = advance_to_run_containing_start(start, end)
    end = split_off_prefix(r, start, end)

    # --- if run is longer than isolation-range we need to split-off a suffix run ---
    if len(r.text) > end:
        split_off_suffix(r, end)
    # --- if run is shorter than isolation-range we need to lengthen it by taking text
    # --- from subsequent runs
    elif len(r.text) < end:
        lengthen_run(r, r_idx, end)

    return Run(r, paragraph)


EXAM_PAPERS = "2021_cpp_papers.docx" # "BiletyIVT.docx" #"2021_java_papers.docx" #  #" "2021_java_papers.docx"

PATTERN_EX_BETW = r"Екзамінаційний білет № (\d+) (\s+) Екзамінаційний білет № \d+"

PATTERN_EX = r"Екзамінаційний білет № (\d+)"
PATTERN_DIG = r"(\d+)."
PATTERN_PAR = r"^(\d+).\t"

doc_papers = Document(EXAM_PAPERS)

dict_papers = {}
k = 0
question_indx = 0
tmp = " "
found_header= False

for p in doc_papers.paragraphs:


    for r in p.runs:

        title_search = re.search(PATTERN_EX, r.text)
        # print("paragprahp",r.text)
        # input()
        if title_search:
            print(r.text)
            dict_papers[k] = tmp
            number = title_search.group(1)
            print(number)
            k = int(number)
            tmp= r.text + "\n\n"
            found_header = True
            question_indx = 0
        elif found_header and re.search(PATTERN_DIG, r.text):
            d = re.search(PATTERN_DIG, r.text).group(1)
            k = 10*k + int(d)
            # question_indx += 1
            #dict_papers[k] = tmp + r.text
            print("head", r.text)
            input()
            tmp += r.text
            # tmp += "\n\n" + str(question_indx)
            found_header= False
        else:
            if k<=0:
                continue
            found_header= False
            if not r.text:
                continue

            print("ppp", r.text)


            if re.search(PATTERN_PAR, r.text):
                question_indx += 1
                print("dddd", r.text)
                # pass
                tmp += "\n\n" + str(question_indx) +"\t"

            # question_indx += 1
            # tmp += "\n" + str(question_indx) +"."

            tmp += r.text
            print("t", tmp)
            input()


dict_papers[k] = tmp

for x,y in dict_papers.items():
    print("numer ", x)
    print("text: ", y)
    # write to file
    print("2021_cpp_ep_{}.docx".format(x))
    """
    text_file = open("2021_cpp_ep_{}.docx".format(x), "w")
    text_file.write(y)
    text_file.close()
    """
    dname = "2021_cpp_ep_{}.docx".format(x)
    document = Document()
    document.add_paragraph(y)
    document.save(dname)


# email body
"""
Name,
course,
repa,
repa
addres
"""
body0="""
Відомості по курсу {}

Шановний(шановна) {}!
Висилаю координати курсу "{}":
Репозиторій з матеріалами лекцій:
{}

У вкладенні пересилаю задачник з 4 теми.

Поточні оцінки та домашні завдання Гугл-документі:
{}

З повагою, 
Ваш викладач
Віктор Бородін. 

Не відповідайте на цей лист - тут працює робот.
У випадку питань пишить на viktorborodinmm@gmail.com

Please don't reply: this message is sent by robot.
In case of any issue reply to viktorborodinmm@gmail.com
"""
# name, repa_lect, title, HW, mark, doc
body1 = """
{} Homework 
Шановний(шановна) {}!
Репозиторій з матеріалами лекцій:
{}
У вкладенні пересилаю задачник з 5 теми.

Ваші домашні завдання з курсу "{}" {}

Ваші поточні бали {}.

Докладніше на Google-docs {}.

Віктор Бородін

Не відповідайте на цей лист - тут працює робот.
У випадку питань пишить на viktorborodinmm@gmail.com

Please don't reply: this message is sent by robot.
In case of any issue reply to viktorborodinmm@gmail.com
"""
# name, title, num, text
body_exst = """
ЦЕ _ ТЕСТ СИСТЕМИ!!!!
   
Перевести сообщение
Отключить для языка: английский

Шановний(шановна){}!
Ваш білет з курсу "{}" номер {} у вкладенні.
Текст білету продубльований сюди:
{} 
Якщо погано відображається - то дивиться у вкладення.

У випадку запитань або проблем з відкриттям 
звертайтеся до викладача на конференцію або телеграм. 

Віктор Бородін

Не відповідайте на цей лист - тут працює робот.
У випадку питань пишить на viktorborodinmm@gmail.com

Please don't reply: this message is sent by robot.
In case of any issue reply to viktorborodinmm@gmail.com

"""

body_ex_end ="""
Шановний(шановна){}!
Ваша оцінка за іспит з курсу "{}" {} та ітогова оцінка {}.
Якщо ви згодні з цією оцінкою пришлить відповіь на viktorborodinmm@gmail.com
або телеграм.
 
Інакше, чекою вас на продовженні іспиту в домовлений час.

Віктор Бородін 

Довідкова інформація на Google-docs {}

Не відповідайте на цей лист - тут працює робот.
У випадку питань пишить на viktorborodinmm@gmail.com

Please don't reply: this message is sent by robot.
In case of any issue reply to viktorborodinmm@gmail.com
"""



EMAIL = 'vborodin001@gmail.com'
PASSWORD = 'iaGM7ahwae'
SERVER = 'imap.gmail.com'

email_id='vborodin001@gmail.com'
email_pass='iaGM7ahwae'

SMTP_HOST = 'smtp.gmail.com'
SMTP_PORT = 465

from email.header import Header
from email.mime.base import MIMEBase
from email import encoders
from email.mime.multipart import MIMEMultipart
from pathlib import Path

from email.mime import application

REPA_LECT = students_data[KURS]["repa"]

COURSE = students_data[KURS]["title"]
REPA = students_data[KURS]["groups"][0]["repa"]
ADDRESS_GD = students_data[KURS]["groups"][0]["doc"]

print()

for k, to in enumerate(DEBUG_TO): # enumerate(lst_to): #

    # msg = EmailMessage()
    # msg['Subject'] = "C++ Course" if KURS == "Cpp" else 'Java Course'
    # msg['From'] = email_id
    # msg['To'] = i

    name = lst_names[k]

    # base
    #body = MIMEText(body0.format(name, COURSE, REPA, ADDRESS_GD))

    # hw
    # body = body1.format(hw[k],lst_marks[k], ADDRESS_GD)

    # exam start
    # random
    num = randint(1, NUM_STUD)
    # or chosen
    # num = lst_exam_papers[k]
    z = dict_papers[num]
    # print(body) # body.format(r1,r2,)
    print("{}: to {} for {} ".format(k, i, name))

    """ 
    msg.set_content(body)  # if you want to add an attachment
    # files= ['lectС.pdf', 'methodicalCPP.pdf']   #
    files = ['LabsJava1_2.pdf']  # ["2021_java_ep_{}.docx".format(num)]


    for file in files:
        with open(file, 'rb') as f:
            data = f.read()
            name = f.name
            msg.add_attachment(data,
                               maintype='application',
                               subtype='octet-stream',
                               filename=name)
    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
        smtp.login(email_id, email_pass)
        print("login done")
        smtp.send_message(msg)
        # s.sendmail(sender, recipients.split(','), msg.as_string())
        smtp.quit()
        print("messages sent")
    """
    msg = MIMEMultipart()
    if MODE==0:
        msg.attach(MIMEText(body0.format(KURS,name, COURSE, REPA, ADDRESS_GD), 'plain', 'utf-8'))
    elif MODE==1:
        # name, repa_lect, title, HW, mark, doc
        Text = body1.format(KURS, name, students_data[KURS]["repa"], COURSE,lst_hw[k],lst_marks[k],ADDRESS_GD
                            )
        print(Text)

        msg.attach(MIMEText(Text, 'plain', 'utf-8'))
    elif MODE==2:
        Text = body_exst.format(name, COURSE, num, z)
        print(Text)

        msg.attach(MIMEText(Text, 'plain', 'utf-8'))

    msg['Subject'] = Header("C++ Course" if KURS == "C++" else 'Java Course', 'utf-8')
    msg['From'] = email_id
    msg['To'] = to  #", ".join(recipients_emails)

    if MODE==1:
        files = ['LabsJava5.pdf'] #students_data[KURS]["man_books"] # ,  ['lectС.pdf', 'methodicalCPP.pdf']  #
        # for attach in files:
        #
        #     attach_file = open(attach, 'rb')  # Open the file as binary mode
        #     payload = MIMEBase('application', 'octate-stream')
        #     payload.set_payload((attach_file).read())
        #     encoders.encode_base64(payload)  # encode the attachment
        #     # add payload header with filename
        #     payload.add_header('Content-Decomposition', 'attachment', filename=attach)
        #     msg.attach(payload)


        for path in files:
            #path = os.path.join("~/Лекции/Cpp/CppCourse","methodicalCPP.pdf")
            part = MIMEBase('application', "octet-stream")
            with open(path, 'rb') as file:
                part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition',
                            'attachment; filename={}'.format(Path(path).name), filename = Path(path).name)
            print("attached: ",Path(path).name)
            msg.attach(part)
    elif MODE==2:
        files = ["2021_cpp_ep_{}.docx".format(num)]


        for path in files:
            #path = os.path.join("~/Лекции/Cpp/CppCourse","methodicalCPP.pdf")
            part = MIMEBase('application', "octet-stream")
            with open(path, 'rb') as file:
                part.set_payload(file.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition',
                            'attachment; filename={}'.format(Path(path).name), filename = Path(path).name)
            print("attached: ",Path(path).name)
            msg.attach(part)


    s = smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT)
    # s =smtplib.SMTP('smtp.gmail.com', 465, timeout=10) # 587
    # s.set_debuglevel(1)
    try:
        #s.starttls()
        s.login(EMAIL, PASSWORD)
        s.sendmail(msg['From'], to, msg.as_string())
    finally:
        s.quit()
        print("All sent")